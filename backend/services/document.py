import io

from core.exceptions import DocumentLinkedError
import llm.rag as rag_service
from db.models import Document
from docx import Document as WordDocument
from fastapi import File, UploadFile
from pypdf import PdfReader
from schemas.document import DocumentCreatePaste
from sqlalchemy.orm import Session
from utils.files import validate_file, validate_text_length

from services.base import BaseService


class DocumentService(BaseService[Document]):
    def __init__(self, session: Session) -> None:
        super().__init__(session, Document)

    @staticmethod
    def _process_pdf_file(content: bytes):
        pdf_stream = io.BytesIO(content)

        reader = PdfReader(pdf_stream)

        pages_texts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text().strip()

            if page_text:
                page_text += f"\n -- Page {i + 1} --- \n"
                pages_texts.append(page_text)

        total_text = "".join(pages_texts)

        return total_text

    @staticmethod
    def _process_docx_file(content: bytes):
        stream = io.BytesIO(content)
        doc = WordDocument(stream)

        doc_texts = []
        for para in doc.paragraphs:
            para_text = para.text.strip()

            if para_text:
                doc_texts.append(para_text)

        for table in doc.tables:
            for row in table.rows:
                cells_texts = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if cells_texts:
                    doc_texts.append(" | ".join(cells_texts))

        total_text = "\n\n".join(doc_texts)

        return total_text

    @staticmethod
    def _process_txt_file(content: bytes):
        text = content.decode(errors="ignore")

        return text.strip()

    def _process_file(self, file: UploadFile = File(...)):
        if not file.filename:
            raise ValueError("File must have a name")

        ext = validate_file(file.filename)

        content = file.file.read()

        if ext == "pdf":
            return (ext, self._process_pdf_file(content))
        elif ext == "docx":
            return (ext, self._process_docx_file(content))

        return (ext, self._process_txt_file(content))

    def create_document_by_paste(self, doc_create_paste: DocumentCreatePaste):
        validate_text_length(doc_create_paste.content)
        document = Document(
            **doc_create_paste.model_dump(), char_count=len(doc_create_paste.content)
        )

        self.session.add(document)
        self.session.flush()

        rag_service.save_document_content(document.id, document.content)

        self.session.commit()

        return document

    def create_document_by_upload(self, file: UploadFile):
        format, text = self._process_file(file)

        validate_text_length(text)

        document = Document(
            name=file.filename,
            source_type="upload",
            format=format,
            content=text,
            char_count=len(text),
        )
        self.session.add(document)
        self.session.flush()
        rag_service.save_document_content(document.id, document.content)
        self.session.commit()

        return document

    def delete_document(self, document_id: int):
        document = self.get_or_404(document_id)

        if document.quiz_generations:
            raise DocumentLinkedError("The document is already being used in a quiz")

        self.session.delete(document)
        rag_service.delete_document_content(document_id)

        self.session.commit()
